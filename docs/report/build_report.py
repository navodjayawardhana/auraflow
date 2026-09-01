# -*- coding: utf-8 -*-
"""Build the CMP 7003 PRAC1 draft report as a Word document.

Formatting follows the assessment brief: A4, Calibri 11, 1" LHS/RHS margins,
0.5" binding gutter, 1" header and footer.
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"E:\MSC\Emerging Mobile Applications\auraflow\docs\report\AuraFlow_CMP7003_PRAC1_DRAFT-v1.docx"

doc = Document()

# ---------------------------------------------------------------- page setup
sec = doc.sections[0]
sec.page_width = Cm(21.0)
sec.page_height = Cm(29.7)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)
sec.header_distance = Inches(0.5)
sec.footer_distance = Inches(0.5)
sec.gutter = Inches(0.5)

# ---------------------------------------------------------------- base styles
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.15

for name, size, colour, before in (
    ('Heading 1', 16, 0x1F3864, 18),
    ('Heading 2', 13, 0x2E5496, 12),
    ('Heading 3', 11.5, 0x2E5496, 10),
):
    st = doc.styles[name]
    st.font.name = 'Calibri'
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor(colour >> 16, (colour >> 8) & 0xFF, colour & 0xFF)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.keep_with_next = True

cap = doc.styles.add_style('Caption Text', 1)
cap.font.name = 'Calibri'
cap.font.size = Pt(9.5)
cap.font.bold = True
cap.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
cap.paragraph_format.space_before = Pt(6)
cap.paragraph_format.space_after = Pt(4)
cap.paragraph_format.keep_with_next = True

mono = doc.styles.add_style('Formula', 1)
mono.font.name = 'Consolas'
mono.font.size = Pt(10)
mono.paragraph_format.left_indent = Cm(1.0)
mono.paragraph_format.space_before = Pt(4)
mono.paragraph_format.space_after = Pt(8)

# ---------------------------------------------------------------- footer page no
def add_page_numbers(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.style = doc.styles['Normal']
    for run in list(p.runs):
        run._element.getparent().remove(run._element)
    r = p.add_run()
    r.font.size = Pt(9)
    for instr in ('begin', 'PAGE', 'end'):
        el = OxmlElement('w:fldChar' if instr in ('begin', 'end') else 'w:instrText')
        if instr in ('begin', 'end'):
            el.set(qn('w:fldCharType'), instr)
        else:
            el.set(qn('xml:space'), 'preserve')
            el.text = ' PAGE '
        r._r.append(el)

add_page_numbers(sec)

# ---------------------------------------------------------------- rich text
TOKEN = re.compile(r'(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`|\u27e6.+?\u27e7|\u3010.+?\u3011)', re.S)

def write_runs(par, text):
    """Render **bold**, *italic*, `code`, and highlight editorial markers."""
    for piece in TOKEN.split(text):
        if not piece:
            continue
        if piece.startswith('**'):
            par.add_run(piece[2:-2]).bold = True
        elif piece.startswith('*'):
            par.add_run(piece[1:-1]).italic = True
        elif piece.startswith('`'):
            r = par.add_run(piece[1:-1])
            r.font.name = 'Consolas'
            r.font.size = Pt(10)
        elif piece.startswith('\u27e6') or piece.startswith('\u3010'):
            from docx.enum.text import WD_COLOR_INDEX
            r = par.add_run(piece)
            r.font.highlight_color = WD_COLOR_INDEX.YELLOW
            r.bold = True
        else:
            par.add_run(piece)

def para(text='', style=None, align=None, space_after=None):
    p = doc.add_paragraph(style=style)
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    write_runs(p, text)
    return p

def h(level, text):
    p = doc.add_heading(level=level)
    write_runs(p, text)
    return p

def bullets(items, numbered=False):
    for it in items:
        p = doc.add_paragraph(style='List Number' if numbered else 'List Bullet')
        p.paragraph_format.space_after = Pt(4)
        write_runs(p, it)

def table(headers, rows, caption=None, widths=None):
    if caption:
        para(caption, style='Caption Text')
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, txt in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        write_runs(p, txt)
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9.5)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'D9E2F3')
        hdr[i]._tc.get_or_add_tcPr().append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, txt in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            write_runs(p, str(txt))
            for r in p.runs:
                r.font.size = Pt(9.5)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

FIG_DIR = Path(__file__).resolve().parent / 'figures'

def figure(path, width_cm, caption):
    """Embed a generated figure. Falls back to a slot if the PNG is missing."""
    if not Path(path).exists():
        figure_slot(caption, 'missing ' + Path(path).name + ' \u2014 run its build script')
        return
    para(caption, style='Caption Text')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    p.add_run().add_picture(str(path), width=Cm(width_cm))

def figure_slot(caption, hint):
    para(caption, style='Caption Text')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    from docx.enum.text import WD_COLOR_INDEX
    r = p.add_run('\u27e6 INSERT FIGURE \u2014 ' + hint + ' \u27e7')
    r.bold = True
    r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    p.paragraph_format.space_after = Pt(10)

# ================================================================ TITLE PAGE
t = para('AuraFlow', align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
t.runs[0].font.size = Pt(30)
t.runs[0].bold = True
t.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

s = para('An AI-Driven Smart Lifestyle Companion', align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
s.runs[0].font.size = Pt(15)
s.runs[0].font.color.rgb = RGBColor(0x40, 0x40, 0x40)

for line, sz, bold in (
    ('CMP 7003 \u2014 Emerging Mobile Applications', 12, True),
    ('PRAC1 \u2014 Practical Project', 12, False),
    ('Cardiff Metropolitan University / ICBT', 11, False),
    ('Academic Year 2025\u20132026, Semester 2', 11, False),
):
    p = para(line, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    p.runs[0].font.size = Pt(sz)
    p.runs[0].bold = bold

doc.add_paragraph().paragraph_format.space_after = Pt(18)

table(
    ['', ''],
    [['Name', '\u27e6 your name \u27e7'],
     ['Cardiff Met ID', '\u27e6 \u27e7'],
     ['ICBT ID', '\u27e6 \u27e7'],
     ['Batch Number', '\u27e6 \u27e7'],
     ['Word count', '\u27e6 fill after final trim \u2014 target 3,000 \u27e7'],
     ['Submission date', '\u27e6 \u27e7']],
    widths=[5.0, 9.5],
)

doc.add_paragraph().paragraph_format.space_after = Pt(14)
p = para('DRAFT v1 \u2014 not for submission. Every highlighted marker must be resolved '
         'and removed, every citation marked VERIFY must be confirmed against the source, '
         'and the text must be revised into the author\u2019s own voice before submission. '
         'Declare the use of AI assistance in Appendix G.',
         align=WD_ALIGN_PARAGRAPH.CENTER)
p.runs[0].font.size = Pt(9.5)
p.runs[0].italic = True
p.runs[0].font.color.rgb = RGBColor(0x99, 0x33, 0x00)

doc.add_page_break()

# ================================================================ CONTENTS
h(1, 'Contents')
p = doc.add_paragraph()
r = p.add_run()
fld = OxmlElement('w:fldChar'); fld.set(qn('w:fldCharType'), 'begin'); r._r.append(fld)
it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve')
it.text = ' TOC \\o "1-3" \\h \\z \\u '
r._r.append(it)
fld = OxmlElement('w:fldChar'); fld.set(qn('w:fldCharType'), 'separate'); r._r.append(fld)
r2 = p.add_run('\u27e6 In Word: right-click here \u2192 Update Field \u2192 Update entire table \u27e7')
r2.italic = True
r3 = p.add_run()
fld = OxmlElement('w:fldChar'); fld.set(qn('w:fldCharType'), 'end'); r3._r.append(fld)

doc.add_page_break()

# ================================================================ 1
h(1, '1. Introduction')

h(2, '1.1 Problem and motivation')
para('Knowledge work is scheduled by the clock. Calendars allocate demanding tasks to '
     'whatever hour is free rather than to the hours in which a person is physiologically able '
     'to do them, a mismatch recognised as a source of lost productivity and strain '
     '(Roenneberg \u2018et al.\u2019, 2003)\u3010VERIFY\u3011. Consumer wearables now measure the signals that '
     'plausibly indicate readiness \u2014 sleep architecture, resting heart rate, heart-rate '
     'variability \u2014 continuously and unobtrusively.')
para('Two obstacles separate that measurement from a useful recommendation. Structurally, '
     'wearable data sits behind proprietary protocols and vendors do not disclose how their '
     'readiness scores are derived. Epistemically, those scores are presented with a precision '
     'their accuracy does not support, and none of that accuracy is published. AuraFlow '
     'addresses both: it recommends **when** to schedule demanding work, computes that '
     'recommendation on the device from an interpretable model, and discloses how much of it '
     'rests on the user\u2019s own data rather than a population median.')

h(2, '1.2 Aim and objectives')
para('**Aim:** to design, implement and critically evaluate a mobile lifestyle companion that '
     'recommends cognitively favourable working hours from wearable and contextual signals, '
     'and to establish whether such a recommendation is achievable at consumer-wearable '
     'fidelity.')
table(['', 'Objective'],
      [['**O1**', 'Compare existing intelligent lifestyle systems to establish a design gap.'],
       ['**O2**', 'Design a modular, secure, scalable architecture separating client from backend.'],
       ['**O3**', 'Implement a cross-platform application integrating on-device machine learning, a server-side language model, wearable and IoT sensing, and computer vision.'],
       ['**O4**', 'Evaluate predictive performance against explicit baselines under participant-wise validation.'],
       ['**O5**', 'Evaluate usability, performance and security, and report limitations without overstatement.']],
      widths=[1.6, 12.9])

h(2, '1.3 Contributions and report structure')
para('The report contributes measured primary evidence of vendor lock-in in a consumer '
     'wearable, obtained by protocol inspection rather than cited (\u00a73.3); the finding that '
     'wearable biometrics add almost nothing to hour-level readiness prediction over '
     'time-of-day alone, against the project\u2019s own founding premise (\u00a75.4); and an '
     'architecture that discloses that weakness rather than concealing it (\u00a74.4). Section 2 '
     'reviews related systems and theory, Section 3 the methodology, Section 4 design and '
     'implementation, Section 5 the evaluation, and Section 6 concludes.')

# ================================================================ 2
h(1, '2. Literature Review')

h(2, '2.1 Context-aware and intelligent mobile systems')
para('Context-aware computing was formalised by Dey (2001) and Abowd \u2018et al.\u2019 (1999), and '
     'mobile platforms have since made its sensing substrate ubiquitous. Three shifts changed '
     'what a handset can do: on-device inference removed the network round trip; large '
     'language models moved assistance from scripted intent matching to open dialogue; and '
     'wearable sensing extended context from the device\u2019s environment to the user\u2019s '
     'physiology. The result is applications that predict capability rather than merely record '
     'behaviour \u2014 importing a burden the recording generation did not carry, since a '
     'prediction can be wrong and a system that predicts owes its user an account of how '
     'often.')

h(2, '2.2 Critical comparison of existing systems')
table(['System', 'Intelligence', 'On device', 'Publishes accuracy', 'Discloses provenance'],
      [['Whoop', 'Proprietary recovery model', 'No', 'No', 'No'],
       ['Oura', 'Proprietary readiness score', 'Partial', 'No', 'No'],
       ['Fitbit Premium', 'Sleep and readiness scores', 'No', 'No', 'No'],
       ['Rise Science', 'Circadian sleep-debt model', 'No', 'No', 'No'],
       ['Reclaim.ai', 'Calendar heuristics only', 'No', 'n/a', 'n/a'],
       ['Samsung Health', 'Aggregation and scores', 'Partial', 'No', 'No'],
       ['**AuraFlow**', '**Interpretable on-device model**', '**Yes**', '**Yes (\u00a75.4)**', '**Yes (\u00a74.4)**']],
      caption='Table 1 \u2014 Comparison of existing intelligent lifestyle systems',
      widths=[2.8, 4.4, 1.9, 2.6, 2.8])
para('Two patterns emerge. Biometric systems produce a readiness figure but stop short of '
     'acting on it, and none reports the predictive validity of the score it sells; '
     'scheduling systems act, but on calendar heuristics with no physiological input. No '
     'system closes the loop from measured physiology to a scheduling decision while '
     'disclosing how reliable that decision is.')

h(2, '2.3 Theoretical foundations')
para('**Russell\u2019s (1980) circumplex model** decomposes affect into valence '
     '(pleasant\u2013unpleasant) and arousal (activated\u2013deactivated), predicting which subjective '
     'states a wearable can reach: heart rate, HRV, sleep pressure and movement are correlates '
     'of arousal and offer no privileged access to valence. \u00a73.3 uses this to choose the '
     'prediction target; \u00a75.4 tests the prediction. **Borb\u00e9ly\u2019s (1982) two-process model** '
     'describes alertness as a circadian oscillator interacting with homeostatic sleep '
     'pressure, motivating the cyclical hour encoding of \u00a74.3. **Self-Determination Theory** '
     '(Deci and Ryan, 2000) holds autonomy to be a precondition of sustained motivation, so '
     'the system suggests hours rather than rewriting a calendar. **Nielsen\u2019s (1994) '
     'heuristics** \u2014 in particular visibility of system status \u2014 inform the disclosure '
     'interface (\u00a74.4) and supply the instrument of \u00a75.3.')

h(2, '2.4 The gap')
para('Existing systems are opaque in two ways: they do not disclose how much of a '
     'recommendation derives from the individual\u2019s own data rather than population defaults, '
     'and they do not publish the predictive validity of the scores they ask users to act on. '
     'The wearable-validity literature documents substantial measurement error in consumer '
     'devices (Haghayegh \u2018et al.\u2019, 2019)\u3010VERIFY\u3011, yet applications built on them present '
     'their outputs without corresponding uncertainty. AuraFlow closes both gaps, and Section '
     '5 reports its own accuracy in the terms this section finds missing.')

# ================================================================ 3
h(1, '3. Methodology')

h(2, '3.1 Development approach')
para('Development followed an incremental, iterative lifecycle, delivering vertical slices '
     'spanning model, API and interface so that each increment was demonstrable rather than '
     'merely complete at one layer. This suited a project whose central risk \u2014 whether the '
     'prediction target carried any signal at all \u2014 could only be retired by building enough '
     'of the pipeline to measure it. Two practices support the audit trail this report relies '
     'on: eight Architecture Decision Records (Nygard, 2011)\u3010VERIFY\u3011 document each '
     'significant choice and its rejected alternatives, and Conventional Commits across 45 '
     'commits state why each change was made rather than only what. Continuous integration '
     'runs the full test suite on every push.')

h(2, '3.2 Technology stack and justification')
table(['Layer', 'Selected', 'Rationale', 'Rejected'],
      [['Mobile', 'React Native (Expo 54), TypeScript',
        'One codebase, two platforms; typing spans the model port',
        'Flutter \u2014 thinner ML/BLE ecosystem'],
       ['Backend', 'Laravel 13, Sanctum', 'Token auth, queues, validation; layering suits DDD',
        'Node/Express \u2014 more assembly'],
       ['Database', 'SQLite', 'Prototype scale; zero-config for an examiner',
        'Firebase \u2014 the lock-in this project critiques'],
       ['Inference', 'TypeScript port of the pipeline', 'Arithmetic only; no runtime, no download',
        'TFLite \u2014 needless for a linear model'],
       ['LLM', 'Gemini, server-side', 'The key never reaches a user device',
        'On-device LLM \u2014 infeasible']],
      caption='Table 2 \u2014 Technology selection and justification',
      widths=[1.9, 3.4, 4.8, 4.4])

h(2, '3.3 Data strategy and research ethics')
para('The project\u2019s original premise was to train on the author\u2019s own wearable data. Protocol '
     'inspection of the device (Huawei Watch Fit) established that this was impossible: BLE '
     'service discovery found the mandatory Generic Access and Device Information services '
     'and a proprietary, authenticated vendor service (`0xFE86`), but the standard Heart Rate '
     'service `0x180D` was **absent** \u2014 the device does not expose its sensors to a '
     'third-party client. This is reported as a finding rather than an obstacle: it is '
     'directly measured primary evidence of the vendor lock-in identified in \u00a72.4 (Appendix C).')
para('The project therefore uses two published cohorts. **LifeSnaps** (Yfantidou \u2018et al.\u2019, '
     '2022) \u2014 71 participants, four months, CC BY 4.0 \u2014 is the primary training and '
     'evaluation set; **PMData** (Thambawita \u2018et al.\u2019, 2020) \u2014 16 participants, CC BY-NC 4.0 '
     '\u2014 is used for validation only. Integrity was verified against the published MD5 '
     'checksum, and each retrieval is pinned by SHA-256 and timestamp in a provenance '
     'manifest. Because PMData is non-commercial, any commercial pathway rests on LifeSnaps '
     'alone. Both cohorts are de-identified and consented for research reuse.')
para('A seeded simulator supplies demonstration data, task-density values absent from both '
     'cohorts, and edge-case fixtures. **Simulator output never enters model training.** Had '
     'it done so, the model would have learned the simulator\u2019s own rules and every metric in '
     '\u00a75 would be circular, with no error raised anywhere. The rule is enforced in code, not '
     'by convention: data origin is a column, an assertion rejects any frame containing '
     'synthetic or untagged rows, and a test suite fails the build if the boundary is breached.')

h(2, '3.4 Evaluation design')
para('Three axes were defined before any result was available: **E1**, retrospective policy '
     'evaluation measuring whether acting on the recommendation would have helped; **E2**, '
     'usability via the System Usability Scale and a heuristic walkthrough; and **E3**, '
     'validation of the daily Recovery Score against self-reported readiness.')

# ================================================================ 4
h(1, '4. Design and Implementation')

h(2, '4.1 System architecture')
figure_slot('Figure 1 \u2014 System architecture',
            'docs/diagrams/02-architecture.jpg')
para('The system separates a React Native client from a stateless Laravel API, with a Python '
     'training pipeline that produces a deployment artefact rather than a running service. The '
     'API is versioned from its first endpoint (`/v1`): retrofitting a prefix once a client is '
     'in users\u2019 hands means supporting both indefinitely, since app updates cannot be forced.')
para('The backend follows a Domain-Driven Design layering \u2014 Domain, Application, '
     'Infrastructure, Http (Evans, 2003) \u2014 across seven bounded contexts. The layering is '
     'applied **asymmetrically, and deliberately so.** The Wellbeing aggregate, where recovery '
     'scoring embeds genuine domain rules and invariants, uses the full stack of value '
     'objects, domain services, repository interfaces and use cases. Append-only records such '
     'as meals use a thinner slice, because an aggregate wrapped around a record that is '
     'created and never mutated adds indirection without protecting an invariant; uniform '
     'application of a heavyweight pattern to trivial data is a common failure of DDD '
     'adoption. The client mirrors the separation, keeping business logic out of components \u2014 '
     'which is what makes the 98-test mobile suite possible.')

h(2, '4.2 Design patterns')
table(['Pattern', 'Applied at', 'What it bought'],
      [['Adapter', 'HealthProvider \u2192 Health Connect, HealthKit, Replay',
        'Deterministic cohort replay; platform APIs swapped without touching callers'],
       ['Repository', 'Domain interfaces, Eloquent implementations',
        'Domain independent of persistence; use cases tested without a database'],
       ['Strategy', 'Vitals source \u2014 BLE preferred, MQTT fallback',
        'Transport chosen at runtime, one reading in the UI (ADR-0007)'],
       ['Facade', 'GeminiClient; weather provider chain',
        'Provider failure and fallback ordering at one boundary'],
       ['Observer', 'Reactive context providers', 'Screens re-render from state, without polling'],
       ['Value Object', 'RecoveryScore, BodyMassIndex, ResetCode',
        'Validity enforced at construction \u2014 an invalid score cannot exist']],
      caption='Table 3 \u2014 Design patterns applied, and what each one bought',
      widths=[2.3, 5.6, 6.6])
para('The Adapter\u2019s third implementation exists because \u00a73.3 removed live wearable sync, and '
     'converts that loss into a benefit: demonstrations become deterministic rather than '
     'dependent on a device pairing successfully.')

h(2, '4.3 The intelligence split')
para('The system uses three kinds of machine intelligence, each chosen for the shape of its '
     'problem rather than for novelty.')
para('**Numbers.** Readiness prediction is a logistic regression trained in scikit-learn, '
     'exported and re-implemented in TypeScript. It runs on the device as arithmetic \u2014 no '
     'inference runtime, no model download, no network call. Crucially the **entire pipeline** '
     'is exported, not only the coefficients: imputer medians, scaler parameters, coefficients '
     'and intercept. Without the transform parameters the device cannot reproduce the space it '
     'was fitted in, and would silently score plausible nonsense.')
para('z = \u03a3 coefficients[i] \u00d7 (x[i] \u2212 mean[i]) / std[i] + intercept', style='Formula')
para('p = 1 / (1 + exp(\u2212z))', style='Formula')
para('Following Borb\u00e9ly (1982), the hour feature is encoded cyclically as sine and cosine '
     'rather than as an integer, because a linear model treats 23:00 and 00:00 as twenty-three '
     'units apart and cannot otherwise fit a rhythm.')
para('**Words.** The daily brief and assistant call Gemini server-side, so the API key never '
     'reaches a device. Guardrails live in the prompt and the prompt is itself under test '
     '(ADR-0005), and the assistant answers from a grounding pack built from the user\u2019s own '
     'fortnight rather than open generation (ADR-0008).')
para('**Pixels.** The movement coach runs a pre-trained pose estimator on device; frames never '
     'leave it. **The rep counter is not a model:** seventeen landmarks feed joint-angle '
     'trigonometry and a hysteresis state machine, which makes it unit-testable against golden '
     'angle sequences and explainable to a user.')

h(2, '4.4 User interface and user experience design')
figure_slot('Figure 2 \u2014 Interface composite: Today, Insights, Assistant, Movement, Plan, Disclosure',
            'six screenshots at consistent scale, labelled (a)\u2013(f)')
para('The client presents 22 screens across an authentication group and an application group, '
     'with the primary journey \u2014 Today, Insights, Assistant, Movement \u2014 reachable within one '
     'navigation level. Depth was constrained deliberately: a recommendation that takes four '
     'taps to reach will not be consulted in the moment it applies. A token-based design '
     'system keeps spacing, typography and colour on one scale across every screen.')
para('**Accessibility** was treated as a requirement rather than a refinement: text and '
     'background pairs meet WCAG AA contrast, interactive targets meet the 44-point minimum, '
     'controls carry labels and roles for screen readers, type scales with the system '
     'font-size setting, and no state is signalled by colour alone.')
para('The most distinctive element follows from \u00a72.4. The disclosure view tells the user '
     '**how many of the model\u2019s 25 inputs are genuinely theirs** rather than a training-set '
     'median, and the Recovery Score is marked provisional on days lacking a '
     'resting-heart-rate baseline instead of being silently renormalised over whichever '
     'components are present. This implements Nielsen\u2019s visibility of system status (1994) '
     'against the commercial norm of Table 1: a system that knows its own prediction is weak '
     'should say so where the user can see it.')

h(2, '4.5 Security, performance and scalability by design')
para('Authentication uses Sanctum bearer tokens, held in the platform secure store while '
     'cached health data sits in general storage (ADR-0002). Every authenticated route is '
     'scoped to the requesting user and accepts no resource identifier, so no path exists by '
     'which one account can address another\u2019s data. Public endpoints carry two limiters, '
     'per-IP and per-email-and-IP, because email alone would let an attacker lock a known user '
     'out of their own account recovery while IP alone would let one host work through a list '
     'of addresses. Secrets are excluded from version control and verified after every push.')
para('Performance derives from the architecture rather than later optimisation: prediction '
     'runs on device as arithmetic and needs no network, and an offline cache with an outbox '
     'lets the application function and queue writes without connectivity. Expensive work is '
     'queued and throttled. The API holds no session state, so horizontal scaling requires '
     'only a load balancer, and persistence sits behind repository interfaces, so SQLite can '
     'be replaced without touching domain code.')

# ================================================================ 5
h(1, '5. Evaluation')

h(2, '5.1 Functional testing')
table(['Suite', 'Tests', 'Guards against'],
      [['API \u2014 Laravel', '461', 'Contract, authorisation and invariant regressions'],
       ['Mobile \u2014 Jest', '372', 'Service and model-port regressions; state and cache errors'],
       ['ML \u2014 provenance, split', '21', 'Simulated rows in training; participant and temporal leakage'],
       ['ML \u2014 label, baseline, feature', '71', 'Baseline-free reporting; target drift; feature errors'],
       ['**Total**', '**925**', '']],
      caption='Table 4 \u2014 Automated test suites and what each guards against',
      widths=[4.6, 1.6, 8.3])
para('All 925 pass in CI on every push (Appendix C). Golden-vector tests pin the TypeScript '
     'port to figures computed in Python, and `check:model` fails the build if the bundled '
     'coefficients drift. The suite is '
     '**load-bearing**: a test expecting a +10 bpm deviation returned 8.57, exposing a '
     'resting-heart-rate baseline that included the current day. Correcting it raised held-out '
     'recall from 0.449 to 0.536 and cut the largest coefficient from 0.839 to 0.376 \u2014 an '
     'artefact, not a finding.')

h(2, '5.2 Performance and security testing')
table(['Metric', 'Requirement', 'Measured'],
      [['Readiness inference p95', '< 50 ms', '**0.004 ms**'],
       ['Authenticated read, server-side p95', '< 300 ms', '**1.6 ms**'],
       ['Android JS bundle (Hermes)', '\u2014', '**7.4 MB**'],
       ['Cold start on device', '< 3 s', '\u3010TBC\u3011']],
      caption='Table 5 \u2014 Measured non-functional performance',
      widths=[7.0, 3.0, 4.5])
para('Both budgets are met with orders to spare, and both are **asserted in CI**. The latency '
     'figure is server-side by design: `php artisan serve` re-bootstraps PHP per request, '
     'adding ~265 ms a php-fpm deployment does not, so its 418 ms end-to-end p95 measures the '
     'harness, not the API.')
para('Security was tested rather than asserted: **19** tests assert cross-account isolation, '
     '**17** that unauthenticated routes are refused, **5** the login throttles; scanning all '
     '46 commits for keys returned nothing. The design was reviewed against the OWASP Mobile '
     'Top 10 (OWASP, 2024)\u3010VERIFY\u3011 (Appendix C). Two residual risks are stated '
     'rather than hidden: a public MQTT broker (ADR-0003) and foreground-only '
     'location (ADR-0004).')

h(2, '5.3 Usability evaluation')
para('\u27e6 Complete after REPORT-PLAN \xa76 P2. \u27e7 The System Usability Scale (Brooke, '
     '1996) was administered to **n = 5** after four scripted tasks. Mean SUS was '
     '\u3010TBC\u3011 (SD \u3010TBC\u3011) against the acceptability anchor of 68 (Bangor '
     '*et al.*, 2009)\u3010VERIFY\u3011; a Nielsen (1994) walkthrough found '
     '\u3010TBC\u3011 issues, \u3010TBC\u3011 major. \u27e6 Name the two most severe. \u27e7 Five participants detect severe problems but estimate the mean coarsely, so '
     'this is formative evidence.')

h(2, '5.4 Model performance and effectiveness')
para('**Target selection.** The focus self-report is seven mutually exclusive categories '
     'covering 3.15% of hourly rows, so the task is classification. Of four binary targets the **activation** axis scored ROC-AUC 0.654 and '
     '**valence** 0.568 \u2014 as Russell\u2019s (1980) circumplex predicts '
     '(\xa72.3).')
table(['Model', 'ROC-AUC', 'F1', 'Recall'],
      [['**Logistic regression**', '**0.656 \xb1 0.027**', '0.543', '0.503'],
       ['MLP (16, 8)', '0.626 \xb1 0.038', '0.526', '0.490'],
       ['Hour-of-day lookup *(baseline)*', '0.599 \xb1 0.039', '0.443', '0.374'],
       ['Base and personal rate *(baselines)*', '0.500', '0.000', '0.000']],
      caption='Table 6 \u2014 Baselines, participant-wise five-fold CV',
      widths=[6.2, 2.9, 1.9, 1.9])
para('The linear model beats the strongest baseline by 0.057 AUC; **the neural network lost** '
     'on ~2,400 rows. The personal-rate '
     'baseline scoring **exactly** 0.500 is a leakage check that passed: under a '
     'participant-wise split a per-person lookup must collapse to the base rate. Decoding the '
     'cyclical hour coefficients recovers a learned **circadian peak at 09:26** \u2014 '
     'unprompted agreement with chronobiology.')
figure(FIG_DIR / 'fig-model-results.png', 15.5,
       'Figure 3 \u2014 Confusion matrix and strongest coefficients')
table(['Policy', 'P@1', 'P@3', 'P@5'],
      [['AuraFlow *(observed context)*', '0.667', '0.652', '0.653'],
       ['**AuraFlow (time and biometrics)**', '**0.600**', '0.593', '0.578'],
       ['Fixed 09:00 *(\u201cwork in the morning\u201d)*', '0.578', '0.556', '0.542'],
       ['Population hour lookup', '0.533', '0.519', '0.493']],
      caption='Table 7 \u2014 **Effectiveness**: hours ranked by predicted readiness, top k '
              'checked (base rate 0.447)',
      widths=[8.2, 2.1, 2.1, 2.1])
para('**The defensible figure is 0.600, not 0.667**: the strongest features record where a '
     'person *was*, which a deployed system must predict. The model gains +0.153 on a '
     'person\u2019s own base rate but only **+0.022** on \u201cwork in the morning\u201d '
     '\u2014 within noise, so it must not claim to beat generic advice.')
para('**The central finding.** The Recovery Score, validated per participant, reaches '
     'only \u03c1 = 0.123, tying Fitbit\u2019s own sleep score; biometrics add **+0.002** AUC '
     'over time alone against location\u2019s +0.043. Three independent analyses therefore '
     'agree: wearable signals are weak predictors of subjective cognitive state; time and '
     'location carry the model. The disclosure interface of \xa74.4 exists because of it.')

h(2, '5.5 Limitations and threats to validity')
bullets([
    '**Agreement, not causation.** The hours were lived and labelled before being ranked, so only a prospective trial could support a causal claim.',
    '**Recall 0.536.** The model misses 91 of 196 focus-ready hours; a missed good hour is cheaper than a bad recommendation, so the threshold is defensible but stated.',
    '**A counter-intuitive coefficient.** Resting heart rate predicts readiness positively \u2014 likely confounded with time of day, and not interpreted causally.',
    '**Scope.** ROC-AUC 0.656 sits just above the conventional 0.60 floor; both cohorts are Fitbit-based and self-selected; and cross-cohort transfer and end-to-end UI behaviour are untested.',
], numbered=True)

# ================================================================ 6
h(1, '6. Conclusion and Future Work')

h(2, '6.1 Outcomes against objectives')
para('**O1** was met: Table 1 establishes a gap in disclosure and published validity that no '
     'compared system closes. **O2** was met: a seven-context domain-driven backend behind a '
     'versioned, stateless API, layered in proportion to each context\u2019s complexity. **O3** was '
     'met: 22 screens integrating on-device inference, a server-side language model, BLE and '
     'MQTT sensing from a custom ESP32 node, pose estimation and geolocation. **O4** was met '
     'and returned a result contrary to expectation \u2014 the model beats every baseline, but the '
     'margin over trivial advice is small and the biometrics contribute almost none of it. '
     '**O5** was partially met: functional, performance and security evidence is complete, '
     'while the usability study at five participants is formative.')
para('The most valuable outcome is O4\u2019s negative result. A system built to demonstrate that '
     'wearables can schedule cognitive work instead produced measured evidence that, at '
     'consumer-wearable fidelity and this cohort size, they largely cannot \u2014 and was then '
     'designed to disclose that rather than conceal it behind a confident score.')

h(2, '6.2 Future work')
para('**Location prediction** is the highest-value next step, and is quantified rather than '
     'speculative: predicting location from calendar and geofence history would recover the '
     '0.067 P@1 currently available only with observed context, roughly quadrupling the margin '
     'over trivial advice. A **prospective within-subject trial** is the only design that '
     'converts agreement into a causal claim, and is precisely what \u00a73.3 precluded.')
para('Three emerging technologies extend the architecture directly. **Edge computing** is '
     'already partially exploited \u2014 inference and pose estimation run on device \u2014 and its '
     'natural extension is **federated learning**, aggregating per-user coefficient updates '
     'without any biometric leaving the handset, which the exported-pipeline design of \u00a74.3 '
     'reduces to a vector of 25 numbers. **Advanced AI models** offer a route past the linear '
     'ceiling, although \u00a75.4 indicates the binding constraint is cohort size rather than model '
     'capacity. **Extended reality** is the natural successor to the movement coach: the same '
     'landmarks that drive the joint-angle state machine could drive an overlay correcting '
     'form in the user\u2019s field of view.')

# ================================================================ REFS
doc.add_page_break()
h(1, 'References')
para('\u27e6 Harvard, alphabetical. Confirm every entry marked VERIFY against the source before '
     'submission \u2014 a fabricated reference is an academic-integrity finding, not a formatting '
     'error. Target 28\u201335 entries; the list below is the confirmed core and needs roughly '
     'fifteen more from REPORT-PLAN \u00a75. \u27e7')

REFS = [
    "Abowd, G.D., Dey, A.K., Brown, P.J., Davies, N., Smith, M. and Steggles, P. (1999) \u2018Towards a better understanding of context and context-awareness\u2019, *Handheld and Ubiquitous Computing*. Berlin: Springer, pp. 304\u2013307.",
    "Bangor, A., Kortum, P. and Miller, J. (2009) \u2018Determining what individual SUS scores mean: adding an adjective rating scale\u2019, *Journal of Usability Studies*, 4(3), pp. 114\u2013123. \u3010VERIFY\u3011",
    "Borb\u00e9ly, A.A. (1982) \u2018A two process model of sleep regulation\u2019, *Human Neurobiology*, 1(3), pp. 195\u2013204.",
    "Brooke, J. (1996) \u2018SUS: a \u201cquick and dirty\u201d usability scale\u2019, in Jordan, P.W. et al. (eds.) *Usability Evaluation in Industry*. London: Taylor & Francis, pp. 189\u2013194.",
    "Deci, E.L. and Ryan, R.M. (2000) \u2018The \u201cwhat\u201d and \u201cwhy\u201d of goal pursuits: human needs and the self-determination of behavior\u2019, *Psychological Inquiry*, 11(4), pp. 227\u2013268.",
    "Dey, A.K. (2001) \u2018Understanding and using context\u2019, *Personal and Ubiquitous Computing*, 5(1), pp. 4\u20137.",
    "Evans, E. (2003) *Domain-Driven Design: Tackling Complexity in the Heart of Software*. Boston: Addison-Wesley.",
    "Fowler, M. (2002) *Patterns of Enterprise Application Architecture*. Boston: Addison-Wesley.",
    "Gamma, E., Helm, R., Johnson, R. and Vlissides, J. (1994) *Design Patterns: Elements of Reusable Object-Oriented Software*. Reading, MA: Addison-Wesley.",
    "Haghayegh, S., Khoshnevis, S., Smolensky, M.H., Diller, K.R. and Castriotta, R.J. (2019) \u2018Accuracy of wristband Fitbit models in assessing sleep: systematic review and meta-analysis\u2019, *Journal of Medical Internet Research*, 21(11), e16273. \u3010VERIFY\u3011",
    "Martin, R.C. (2017) *Clean Architecture: A Craftsman\u2019s Guide to Software Structure and Design*. Boston: Prentice Hall.",
    "Nielsen, J. (1994) \u2018Enhancing the explanatory power of usability heuristics\u2019, *Proceedings of the SIGCHI Conference on Human Factors in Computing Systems*. New York: ACM, pp. 152\u2013158.",
    "Nygard, M. (2011) *Documenting architecture decisions*. \u3010VERIFY \u2014 non-peer-reviewed; check the module accepts it, and pair with a peer-reviewed source if not\u3011",
    "OWASP (2024) *OWASP Mobile Top 10*. \u3010VERIFY year, URL and accessed date\u3011",
    "Roenneberg, T., Wirz-Justice, A. and Merrow, M. (2003) \u2018Life between clocks: daily temporal patterns of human chronotypes\u2019, *Journal of Biological Rhythms*, 18(1), pp. 80\u201390. \u3010VERIFY\u3011",
    "Russell, J.A. (1980) \u2018A circumplex model of affect\u2019, *Journal of Personality and Social Psychology*, 39(6), pp. 1161\u20131178.",
    "Thambawita, V., Hicks, S.A., Borgli, H. et al. (2020) \u2018PMData: a sports logging dataset\u2019, *Proceedings of the 11th ACM Multimedia Systems Conference*. New York: ACM, pp. 231\u2013236.",
    "Yfantidou, S., Karagianni, C., Efstathiou, S. et al. (2022) \u2018LifeSnaps: a 4-month multi-modal dataset capturing unobtrusive snapshots of our lives in the wild\u2019, *Scientific Data*, 9, 663.",
]
for ref in REFS:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.space_after = Pt(6)
    write_runs(p, ref)

# ================================================================ APPENDICES
doc.add_page_break()
h(1, 'Appendices')
para('Appendices are excluded from the word count \u2014 and from the grade. Nothing that earns a '
     'mark belongs here; only the raw evidence that body-text claims point to.')
table(['', 'Appendix', 'Status'],
      [['A', 'Reproducibility \u2014 commands, seeds, provenance manifest', 'Ready'],
       ['B', 'Simulator generative rules and calibration record', 'Ready'],
       ['C', 'Raw test evidence, BLE service listing, OWASP mapping', '\u3010TBC\u3011'],
       ['D', 'Usability instrument, consent form, raw SUS responses', '\u3010TBC\u3011'],
       ['E', 'Use-case specifications, FR/NFR list, traceability matrix', '\u3010TBC\u3011'],
       ['F', 'Architecture Decision Record index (ADR-0001 \u2026 ADR-0008)', 'Ready'],
       ['G', '**Declaration of AI use** \u2014 required; this draft must be declared', '\u3010TBC\u3011']],
      widths=[1.2, 10.3, 3.0])

para('**Appendix C \u2014 test evidence.** Captured verbatim from the runs; the summary lines '
     'are each suite\u2019s own output rather than retyped text.')
figure(FIG_DIR / 'fig-test-evidence.png', 15.8,
       'Figure C.1 \u2014 Console output of the three suites: 461 + 372 + 92 = 925 tests passing, '
       'with the model-sync build gate confirming the bundled coefficients match the artifact')

doc.save(OUT)

# ------------------------------------------------------------- word count
d2 = Document(OUT)
WORD = re.compile(r"[\w\u00a7\u00b1\u03c1\u00b2\u2013\u2014.%/'-]+")
section, counts, order = 'front matter', {}, []
for child in d2.element.body.iterchildren():
    tag = child.tag.split('}')[-1]
    if tag not in ('p', 'tbl'):
        continue
    text = ' '.join(n.text or '' for n in child.iter() if n.tag.endswith('}t'))
    if tag == 'p':
        style = child.find(qn('w:pPr'))
        style = style.find(qn('w:pStyle')) if style is not None else None
        if style is not None and style.get(qn('w:val')) == 'Heading1':
            section = text.strip()
    if section not in counts:
        counts[section] = 0
        order.append(section)
    counts[section] += len(WORD.findall(text))

BUDGET = {'1. Introduction': 300, '2. Literature Review': 420, '3. Methodology': 380,
          '4. Design and Implementation': 880, '5. Evaluation': 770,
          '6. Conclusion and Future Work': 250}
print('Saved:', OUT)
counted = 0
for name in order:
    n = counts[name]
    if name in BUDGET:
        counted += n
        flag = '' if n <= BUDGET[name] else '  OVER'
        print(f'  {name:34s} {n:5d} / {BUDGET[name]:4d}{flag}')
    else:
        print(f'  {name:34s} {n:5d}   (not counted)')
print(f'  {"COUNTED TOTAL":34s} {counted:5d} / 3000')
